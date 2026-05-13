"""
Convert ALL-CAPTIONS-2026.md to a formatted .docx for Google Docs import.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = Path(__file__).parent.parent / "docs" / "captions" / "ALL-CAPTIONS-2026.md"
OUTPUT = Path(__file__).parent.parent / "outputs" / "LATAM-Independence-Captions-2026.docx"

# Brand colors
NAVY   = RGBColor(0x1a, 0x1a, 0x2e)
GOLD   = RGBColor(0xe8, 0xb8, 0x4b)
GRAY   = RGBColor(0x88, 0x88, 0x88)
BLACK  = RGBColor(0x2c, 0x2c, 0x2c)

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_rule(doc):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DDDDDD")
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_doc(lines):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name  = "Calibri"
    style.font.size  = Pt(11)
    style.font.color.rgb = BLACK

    # ── Cover title ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after  = Pt(6)
    r = title_p.add_run("🌎  Latin America Independence Days 2026")
    set_font(r, size=22, bold=True, color=NAVY)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(2)
    r = sub_p.add_run("Mundo Asia — Bilingual Social Media Captions")
    set_font(r, size=13, color=GOLD, bold=True)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(24)
    r = meta_p.add_run("Español · Tiếng Việt  ·  20 países  ·  2026")
    set_font(r, size=10, color=GRAY, italic=True)

    doc.add_page_break()

    # ── Parse lines ──────────────────────────────────────────────────────────
    i = 0
    current_lang = None  # "es" | "vi" | None

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip horizontal rules and empty metadata header
        if line in ("---", ""):
            i += 1
            continue

        # H1 — document title (skip, already added)
        if line.startswith("# "):
            i += 1
            continue

        # H2 — Month section header (e.g. "## Enero · Tháng 1")
        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(text.upper())
            set_font(r, size=11, bold=True, color=GRAY)
            # underline via border
            add_rule(doc)
            current_lang = None
            i += 1
            continue

        # H3 — Country heading (e.g. "### 🇲🇽 México — 16 de septiembre")
        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            set_font(r, size=14, bold=True, color=NAVY)
            current_lang = None
            i += 1
            continue

        # Language label lines
        if line.startswith("**Español**") or "Español" in line and line.startswith("**"):
            current_lang = "es"
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run("Español")
            set_font(r, size=9, bold=True, color=GOLD)
            i += 1
            continue

        if line.startswith("**Tiếng Việt**") or "Tiếng Việt" in line and line.startswith("**"):
            current_lang = "vi"
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run("Tiếng Việt")
            set_font(r, size=9, bold=True, color=GOLD)
            i += 1
            continue

        # Hashtag lines
        if line.startswith("#") and not line.startswith("##"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line)
            set_font(r, size=9, italic=True, color=GRAY)
            i += 1
            continue

        # Metadata / italics intro line (> ...)
        if line.startswith("> "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run(text)
            set_font(r, size=9, italic=True, color=GRAY)
            i += 1
            continue

        # Regular caption text
        if line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line)
            set_font(r, size=11, color=BLACK)

        i += 1

    return doc


def main():
    text  = INPUT.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc   = build_doc(lines)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
